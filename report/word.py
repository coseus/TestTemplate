# report/word.py
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import base64
import io
from PIL import Image

def generate_docx(findings, project_info):
    doc = Document()

    # === STILURI ===
    style = doc.styles.add_style('Heading1', WD_STYLE_TYPE.PARAGRAPH)
    style.font.size = Pt(16)
    style.font.bold = True

    style = doc.styles.add_style('Heading2', WD_STYLE_TYPE.PARAGRAPH)
    style.font.size = Pt(14)
    style.font.bold = True

    # === COVER ===
    p = doc.add_paragraph("PENTEST REPORT", style='Heading1')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Client: {project_info['client']}", style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Project: {project_info['project']}", style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # === FINDINGS ===
    for i, f in enumerate(findings, 1):
        doc.add_paragraph(f"6.{i} {f.get('id')} - {f.get('title')}", style='Heading2')
        doc.add_paragraph(f"Host: {f.get('host')}")
        doc.add_paragraph(f"CVSS: {f.get('cvss')} | Severity: {f.get('severity')}")
        doc.add_paragraph(f.get('description', ''))
        
        # === CODE ===
        if f.get('code'):
            p = doc.add_paragraph("Proof of Concept:", style='Normal')
            p = doc.add_paragraph(f.get('code'), style='Normal')
            p.style.font.name = 'Courier New'
            p.style.font.size = Pt(9)

        # === IMAGINI ===
        for img_b64 in f.get('images', []):
            try:
                header, b64_data = img_b64.split(',', 1)
                img_bytes = base64.b64decode(b64_data)
                img_stream = io.BytesIO(img_bytes)
                doc.add_picture(img_stream, width=Inches(6))
            except:
                doc.add_paragraph("[Image failed to load]")

        doc.add_paragraph(f"Remediation: {f.get('remediation', '')}")
        doc.add_page_break()

    return doc
